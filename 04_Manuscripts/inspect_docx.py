#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCXファイルの構造を確認するスクリプト
"""

from docx import Document
from pathlib import Path

def inspect_docx(docx_path):
    """DOCXファイルの段落を確認"""
    doc = Document(docx_path)

    print(f"総段落数: {len(doc.paragraphs)}")
    print("\n--- 最初の50段落 ---\n")

    for i, para in enumerate(doc.paragraphs[:50], 1):
        text = para.text.strip()
        if text:
            # スタイル名も表示
            style = para.style.name if para.style else "None"
            print(f"[{i}] [{style}] {text[:100]}")
            if "reference" in text.lower() or text.startswith("1)") or text.startswith("2)"):
                print("  ^^^ 参考文献セクション候補 ^^^")

    print("\n--- 後半50段落 ---\n")
    for i, para in enumerate(doc.paragraphs[-50:], len(doc.paragraphs) - 49):
        text = para.text.strip()
        if text:
            style = para.style.name if para.style else "None"
            print(f"[{i}] [{style}] {text[:100]}")

def main():
    docx_path = Path(__file__).parent / "Manuscript_heatwave_social_isolation.docx"

    if not docx_path.exists():
        print(f"エラー: {docx_path} が見つかりません")
        return

    print("=" * 60)
    print("DOCXファイル構造確認")
    print("=" * 60)
    print(f"ファイル: {docx_path}\n")

    inspect_docx(docx_path)

    print("\n" + "=" * 60)

if __name__ == "__main__":
    main()
