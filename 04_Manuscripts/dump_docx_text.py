#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
DOCXファイルの生テキストをダンプ
"""

from docx import Document
from pathlib import Path

def main():
    docx_path = Path(__file__).parent / "Manuscript_heatwave_social_isolation.docx"
    output_path = Path(__file__).parent / "docx_full_text.txt"

    doc = Document(docx_path)

    with open(output_path, 'w', encoding='utf-8') as f:
        for i, para in enumerate(doc.paragraphs, 1):
            f.write(f"--- Paragraph {i} ---\n")
            f.write(para.text + "\n\n")

    print(f"全テキストを {output_path} に保存しました")
    print(f"総段落数: {len(doc.paragraphs)}")

if __name__ == "__main__":
    main()
